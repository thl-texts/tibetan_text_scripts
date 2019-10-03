'''
Split a volume file into individual texts based on a text start delimiter as defined in the global
variable `text_break_string`. Not ideal. Better to use in combination with milestones and a catalog

'''

from os import listdir, path
from docx import Document
from docx.shared import Pt
from copy import deepcopy

indir = 'in'
outdir = 'out'
filefilterstr = 'KAMA'  # string that begins all files in the IN folder to weed out '.' and '~" files
text_break_string = '༄༅'


def read_in_doc(fnm):
    """
    Read in a Word docx with file name
    :param fnm:
    :return:
    """
    doc = Document(fnm)
    return doc


def copy_doc_to_template(infnm, outfnm):
    """
    Copies unstyled word document paragraphs to a "template" document which is a Word docx with the correct styles
    :param infnm:
    :param outfnm:
    :return:
    """
    indoc = read_in_doc(infnm)
    outdoc = get_template_doc(outfnm)
    pstyl = get_style(outdoc)
    for ind, p in enumerate(indoc.paragraphs):
        outpara = outdoc.add_paragraph(p.text, pstyl)
        do_annotations(outpara, outdoc)

    set_tib_font(outdoc)
    outdoc.save(outfnm)
    return outdoc


def do_annotations(p, mydoc):
    """
    Finds text betwee « and » and applies the Annotations style

    :param p:
    :param mydoc:
    :return:
    """
    annotstyl = get_style(mydoc, "Annotations")
    ptxt = p.text
    pts = ptxt.split("«")
    if len(pts) > 1:
        p.clear()
        p.add_run(pts[0])
        for rn in pts[1:]:
            rnpts = rn.split("»")
            if len(rnpts) == 1:
                print("No closing annotation marker")
                p.add_run(rn, annotstyl)
            else:
                p.add_run(rnpts[0], annotstyl)
                runtxt = rnpts[1]
                if len(rnpts) > 2:
                    print("more than two closing annotation markers!")
                    runtxt = ''.join(rnpts[1:])
                p.add_run(runtxt)



def get_template_doc(fnm):
    """
    Get the template document that has all the styles in it

    :param fnm:
    :return:
    """
    folder = 'resources'
    tplnm = 'tibtext-styled-tpl.docx'
    tplpath = path.join(folder, tplnm)
    doc = Document(tplpath)
    doc.save(fnm)
    return doc


def get_style(doc, stynm="Paragraph"):
    """
    Find a style in the doc by name, defaults to paragraph

    :param doc:
    :param stynm:
    :return:
    """
    for st in doc.styles:
        if st.name == stynm:
            return st


def get_title_p(tmpldoc):
    """
    Returns the paragraph with the text "TITLE HERE" from template

    :param tmpldoc:
    :return:
    """
    for p in tmpldoc.paragraphs:
        ptxt = p.text
        ptxt = ptxt.strip()
        if ptxt == 'TITLE HERE':
            return p
    print("Title paragraph not found!")
    return False


def set_tib_font(doc):
    """
    Set the font in the document to Jomolhari

    :param doc:
    :return:
    """
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.complex_script = True
            r.font.name = "Jomolhari"
            r.font.size = Pt(16)


def add_styles_without_splitting():
    """
    This is the main routine from add_styles2docx.py from which this script is copied

    :return:
    """

    ct = 0
    for fnm in [f for f in listdir(indir) if not f.startswith('.')]:
        print("Processing {}".format(fnm))
        inf = path.join(indir, fnm)
        outf = path.join(outdir, fnm)
        copy_doc_to_template(inf, outf)
        ct += 1

    print("{} documents done!".format(ct))


def split_into_texts(doc, fnmbase, ct):
    """
    Takes a volume or volume chunk document and splits it into individual text documents based on the global
    variable text_break_string which initially is '༄༅' It also looks to see if there are annotation markers « and »
    that are unclosed within the text break and adds the necessary closure/completion

    :param doc:
    :param fnmbase:
    :param ct:
    :return:
    """
    newtexts = []
    curroutdoc = "{}-t{}.docx".format(fnmbase, str(ct).zfill(2))
    newdoc = Document()
    for n, p in enumerate(doc.paragraphs):
        txt = p.text
        # If text break string is not found, add to current document
        if n == 0 or txt.find(text_break_string) == -1:
            newdoc.add_paragraph(txt)
        else:
            ppts = txt.split(text_break_string)
            for idx, ppt in enumerate(ppts):
                if ppt is None or len(ppt) == 0 or ppt == '«' or ppt == '»':
                    continue
                if is_in_annotation(ppt):
                    ppts[idx] = ppts[idx].lstrip('༄༅། ')
                    ppts[idx] += '»'
                    if idx + 1 < len(ppts):
                        ppts[idx + 1] = '«' + ppts[idx + 1].lstrip('༄༅། ')
                # add part to newdoc, save, add to list, and create another new doc
                newdoc.add_paragraph(ppt.lstrip('། »'))
                newdoc.save(path.join(outdir, curroutdoc))
                if idx < len(ppts) - 1 and len(ppt) > 100:
                    newtexts.append(curroutdoc)
                    ct += 1
                    curroutdoc = "{}-t{}.docx".format(fnmbase, str(ct).zfill(2))
                    newdoc = Document()

    # Save the last doc and add it to the list if not there
    newdoc.save(path.join(outdir, curroutdoc))
    if curroutdoc not in newtexts:
        newtexts.append(curroutdoc)
    return newtexts


def is_in_annotation(txtpt):
    """
    Determines if the string given in txtpt ends in the middle of an annotation

    :param txtpt: String
    :return:
    """
    openpos = txtpt.rfind('«')
    closepos = txtpt.rfind('»')
    if openpos > closepos:
        return True
    return False


def process_volume_chunks(coll, ed, vol):
    """
    Iterates through the volume chunks for a certain volume number and splits them into texts with the function of the
    same name. Uses global variable filefilterstr (the prefix for the volume chunk file name) and the volume number to
    filter files in the in folder. For instance, if filefilterstr is "KAMA" and volume is 2, it will only process chunks
    names "KAMA-0002*". The coll and ed parameters are used to name the output text files.

    :param coll: String
        The Collection sigla
    :param ed: String
        The Edition sigla
    :param vol: Int
        The integer number of the volume
    :return: None
    """
    newtexts = []
    newfilebase = "{}-{}-v{}".format(coll, ed, str(vol).zfill(3))
    filter_str = "{}-{}".format(filefilterstr, str(vol).zfill(3))
    filelist = [f for f in listdir(indir) if f.startswith(filter_str)]
    filelist.sort()
    for fnm in filelist:
        ct = len(newtexts) + 1
        indoc = read_in_doc(path.join(indir, fnm))
        print("\tDoing {} ...".format(fnm))
        newtexts = newtexts + split_into_texts(indoc, newfilebase, ct)

    print("There are {} new texts. Now applying annotations!".format(len(newtexts)))
    for ind, fnm in enumerate(newtexts):
        fpth = path.join(outdir, fnm)
        print("\tDoing text {} ({})".format(ind + 1, fpth))
        copy_doc_to_template(fpth, fpth)  # writing over the split text with the text with template and annotations


if __name__ == "__main__":
    # vnum = 1
    for vnum in range(1, 133):
        print("Processing volume {}".format(vnum))
        process_volume_chunks('km', 'pt', vnum)
        print("Volume {} done".format(vnum))


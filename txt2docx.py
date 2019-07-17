from os import listdir
from docx import Document
from docx.shared import Pt


def read_in_text(dir, fnm):
    try:
        with open("{}/{}".format(dir, fnm), 'r') as infl:
            alltxt = infl.readlines()
    except FileNotFoundError as fnfe:
        print(fnfe)
        alltxt = False

    return alltxt


def get_template_doc(fnm):
    tmplfile = 'tibtext-styled-tpl.docx'
    doc = Document(tmplfile)
    doc.save(fnm)
    return doc


def get_paragraph_style(doc):
    for st in doc.styles:
        if st.name == "Paragraph":
            return st


def set_tib_font(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.complex_script = True
            r.font.name = "Jomolhari"
            r.font.size = Pt(16)


if __name__ == "__main__":
    indir = 'in'
    outdir = 'out'
    for fnm in [f for f in listdir(indir) if not f.startswith('.')]:
        print(fnm)
        outfnm = "{}/{}".format(outdir, fnm.replace(".txt", ".docx"))
        mydoc = get_template_doc(outfnm)
        parasty = get_paragraph_style(mydoc)
        txtlines = read_in_text('in', fnm)
        for ln in txtlines:
            mydoc.add_paragraph(ln, parasty)
        set_tib_font(mydoc)
        mydoc.save(outfnm)
        break

    print("done")


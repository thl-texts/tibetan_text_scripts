from os import listdir, mkdir, path, system, getcwd, chdir, remove, path
import shutil
import glob
from docx import Document
from docx.shared import Pt
import argparse


def reinstate_text_files():
    fpath = './workspace/out'
    for f in glob.glob(path.join(fpath, '*.docx')):
        remove(f)
    bakpath = path.join(fpath, 'bak')
    for f in glob.glob(path.join(bakpath, '*.txt')):
        shutil.move(f, fpath)


def summarize_args(kws):
    print("----------- Parameters   ---------")
    print("In directory: {}".format(kws['in']))
    print("Out directory: {}".format(kws['out']))
    if kws['redo']:
        print("Redoing from text files!!!!")
    print("Options:")
    print("\tAdd metadata table: {}".format("yes" if kws['table'] else "no"))
    print("\tDo annotations: {}".format("yes" if kws['annotations'] else "no"))
    print("\tDo Milestones: {}".format("yes" if kws['milestones'] else "no"))


def stardardize_in_files(infld):
    '''
    Standardize the files in the in folder converting them to docx

    :param infld:
    :return:
    '''
    bakfld = path.join(infld, 'bak')
    if not path.exists(bakfld):
        mkdir(bakfld)
    cwd = getcwd()
    for f in listdir(infld):
        if f.endswith('.txt') or f.endswith('.rtf'):
            newf = f.replace('.txt', '-temp.docx').replace('.rtf', '-temp.docx')
            chdir(infld)
            system('textutil -convert docx -output "{}" "{}"'.format(newf, f))
            chdir(cwd)
            shutil.move(path.join(infld, f), path.join(bakfld, f))
        elif f.endswith('.docx') and not f.endswith('-temp.docx'):
            newf = f.replace('.docx', '-temp.docx')
            shutil.move(path.join(infld, f), path.join(infld, newf))


def clean_out_folder(outd):
    for f in listdir(outd):
        if f.endswith('-temp.docx'):
            remove(path.join(outd, f))


def read_in_text(dir, fnm):
    '''
    Reads in plain text to put in Word doc. Original idea, maybe deprecated
    :param dir: String
    :param fnm: String
    :return: List
    '''
    try:
        with open("{}/{}".format(dir, fnm), 'r') as infl:
            alltxt = infl.readlines()
    except FileNotFoundError as fnfe:
        print(fnfe)
        alltxt = False

    return alltxt


def read_in_doc(fnm):
    '''
    Read in a Word docx with file name

    :param fnm: String
    :return: docx.document.Document
    '''
    doc = Document(fnm)
    return doc


def copy_doc_to_template(infnm, outfnm, include_table=True, do_annots=False, mark_miles=True):
    """
    Copies unstyled word document paragraphs to a "template" document which is a Word docx with the correct styles

    :param infnm:
    :param outfnm:
    :param include_table:
    :param do_annots:
    :param mark_miles:
    :return: docx.document.Document
    """

    indoc = read_in_doc(infnm)
    outdoc = get_template_doc(outfnm, include_table)
    pstyl = get_style(outdoc)
    for ind, p in enumerate(indoc.paragraphs):
        outpara = outdoc.add_paragraph(p.text, pstyl)
        if do_annots:
            anntstyl = get_style(outdoc, "Annotations")
            do_annotations(outpara, anntstyl)
        if mark_miles:
            pgst = get_style(outdoc, 'Page Number Print Edition')
            lnst = get_style(outdoc, 'Line Number Print')
            mark_milestones(outpara, pgst, lnst)
    set_tib_font(outdoc)
    outdoc.save(outfnm)
    print("Saved to {}".format(outfnm))
    return outdoc


def do_annotations(p, annotstyl):
    '''
    Finds text between « and » and applies the Annotations style
        «  »
    :param p:
    :param annotstyl:
    :return: Null (manipulation done by reference)
    '''

    print("\tApplying styles to annotations!")

    # Other option for annot_open and _close: < >
    annot_open = '«'
    annot_close = '»'
    newruns = []
    for rn in p.runs:
        rtxt = rn.text
        pts = rtxt.split(annot_open)
        if len(pts) > 1:
            newruns.append((pts[0], ''))
            for subrn in pts[1:]:
                rnpts = subrn.split(annot_close)
                if len(rnpts) == 1:
                    print("No closing annotation marker")
                    newruns.append((subrn, annotstyl))
                else:
                    newruns.append((rnpts[0], annotstyl))
                    runtxt = rnpts[1]
                    if len(rnpts) > 2:
                        print("more than two closing annotation markers!")
                        runtxt = ''.join(rnpts[1:])
                    newruns.append((runtxt, ''))
        else:
            newruns.append((rn.text, rn.style))

    p.clear()
    for rdata in newruns:
        p.add_run(rdata[0], rdata[1])


def mark_milestones(p, pgstyl, lnstyl):
    print("\tApplying styles to milestones!")
    # Iterate through runs in paragraph
    newruns = []
    for rn in p.runs:
        currsty = rn.style      # get run's current style
        rntxt = rn.text         # get run text
        pts = rntxt.split('[')  # split on milestone delimiter [
        if len(pts) > 1:        # if there are more than one parts, it has a bracket/milestone in it
            for msrn in pts:    # iterate through the parts from splitting on [
                if ']' in msrn:                                 # if the part has a close bracket ], then ...
                    rnpts = msrn.split(']')                     # split it on the ]
                    newms = '[' + rnpts[0] + ']'                # rebuild the milestone string
                    sty = lnstyl if '.' in newms else pgstyl    # choose the style based on whether there is a "."
                    newruns.append((newms, sty))                   # add the milestone with the style
                    if len(rnpts) > 1:                          # if there's another part
                        newruns.append((rnpts[1], currsty))           # add it back with the run's original style
                else:
                    newruns.append((msrn, currsty))  # Otherwise no milestone in this part, add back with style
        else:
            newruns.append((rn.text, rn.style))

    p.clear()
    for rdata in newruns:
        p.add_run(rdata[0], rdata[1])


def get_template_doc(fnm, with_table=True):
    '''
    Get the template document that has all the styles in it
    :param fnm:
    :param with_table:
    :return: docx.document.Document
    '''
    folder = 'resources'
    tplnm = 'tibtext-styled-tpl-2023-09-26.docx'  # 'tibtext-styled-tpl.docx'
    # print("Template: {}".format(tplnm))
    tplpath = path.join(folder, tplnm)
    doc = Document(tplpath)
    parastyl = get_style(doc)
    if not with_table:
        for tbl in doc.tables:
            tbl._element.getparent().remove(tbl._element)
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

    doc.save(fnm)
    return doc


def get_style(doc, stynm="Paragraph"):
    '''
    Find a style in the doc by name, defaults to paragraph

    :param doc:
    :param stynm: The name of the style to find, defaults to paragraph
    :return: docx.styles.styles.Styles
    '''
    for st in doc.styles:
        if st.name == stynm:
            return st


def get_title_p(tmpldoc):
    '''
    Returns the paragraph with the text "TITLE HERE" from template

    :param tmpldoc: docx.Document
    :return: docx.text.paragraph.Paragraph
    '''
    for p in tmpldoc.paragraphs:
        ptxt = p.text
        ptxt = ptxt.strip()
        if ptxt == 'TITLE HERE':
            return p
    print("Title paragraph not found!")
    return False


def set_tib_font(doc):
    '''
    Set the font in the document to Jomolhari

    :param doc:
    :return: Null (manipulation by reference)
    '''
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.complex_script = True
            r.font.name = "Jomolhari"
            r.font.size = Pt(16)


def convert_files(ind, outd, table, annots, milestones):
    ct = 0
    for fnm in [f for f in listdir(ind) if f.endswith('.docx') and not f.startswith('~')]:
        print("Processing {}".format(fnm))
        inf = path.join(ind, fnm)
        outf = path.join(outd, fnm.replace('-temp', ''))
        copy_doc_to_template(inf, outf, table, annots, milestones)
        ct += 1

    print("{} documents done!".format(ct))


parser = argparse.ArgumentParser(description='Add convert text files to Word docs with THL Styles')
parser.add_argument('-i', '--in', default='workspace/out',
                    help='Where are the documents to convert?')
parser.add_argument('-o', '--out',
                    help='Where to save the converted documents')
parser.add_argument('-t', '--table', action='store_true',
                    help='Add metadata tables to documents')
parser.add_argument('-a', '--annotations', action='store_true',
                    help='Convert << and >> into annotation style')
parser.add_argument('-m', '--milestones', action='store_true',
                    help='Add styles to page and line milestones')
parser.add_argument('-r', '--redo', action='store_true',
                    help='Reconvert original text files')
args = parser.parse_args()

if __name__ == "__main__":
    '''
    The main routine that runs the conversion. 
    '''
    kwargs = vars(args)
    if not kwargs['out']:
        kwargs['out'] = kwargs['in']
    indir = kwargs['in']
    if not path.isdir(indir):
        print("The In directory setting, {}, is not a directory".format(indir))
        exit(0)
    outdir = kwargs['out']
    if not path.isdir(outdir):
        print("The Out directory setting, {}, is not a directory".format(outdir))
        exit(0)

    summarize_args(kwargs)
    if kwargs['redo']:
        reinstate_text_files()
    stardardize_in_files(indir)
    print("Files converted to docx!")
    convert_files(indir, outdir, kwargs['table'], kwargs['annotations'], kwargs['milestones'])
    print("Removing temp files ...")
    clean_out_folder(outdir)
    print("Done")

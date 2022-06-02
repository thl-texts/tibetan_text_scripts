import os
from tibtexts.ocrvolpgs import OCRVol
from docx import Document
import argparse


def make_doc(v, stpg, enpg, outfl):
    doc = Document('./resources/blank_doc.docx')
    enpg = v.get_last() if enpg is False else enpg
    para = doc.paragraphs[0]
    for pn in range(stpg, enpg + 1):
        pg = v.get_page(pn)
        # para = doc.add_paragraph()
        para.add_run(f"[{pn}]", "Page Number Print Edition")
        for n, ln in enumerate(pg):
            if n == 0:
                ln = ln.lstrip('༄༅། །་༈༑༏་')
            lnum = n + 1
            para.add_run(f"[{pn}.{lnum}]", "Line Number Print")
            para.add_run(ln)
    outfl = f"workspace/out{outfl}"
    doc.save(outfl)


def get_args():
    parser = argparse.ArgumentParser(description='Insert Milestones from OCR in Unicode Docs')
    parser.add_argument('vnum', type=int, help="Number of volume to load")
    parser.add_argument('outfile', type=str, help="Name of outfile to write")
    parser.add_argument('-s', '--start', type=int, default=1,
                        help='The scan page the volume starts at')
    parser.add_argument('-e', '--end', type=int, required=False,
                        help='The scan page the volume end at (inclusive)')
    parser.add_argument('-o', '--offset', type=int, default=0,
                        help='Number of ocr pages before first text page')
    data = parser.parse_args()
    return vars(data)


def main():
    args = get_args()
    volnum = args['vnum']
    outfile = args['outfile']
    offset = args['offset']
    stpg = args['start']
    endpg = args['end']
    vol = OCRVol(volnum, offset)
    vol.load_pages()
    make_doc(vol, stpg, endpg, outfile)


if __name__ == "__main__":
    main()

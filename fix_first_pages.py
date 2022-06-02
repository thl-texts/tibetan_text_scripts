"""
A script to add first page milestones if missing from word document chunks
"""

import os
import argparse
import shutil
from docx import Document
import re
import copy
from docx.text.run import Run
from docx.oxml.text.run import CT_R


parser = argparse.ArgumentParser(description='Add first page milestones to all documents in folder')
parser.add_argument('-f', '--folder', default='workspace/out',
                    help='Folder with documents')
args = parser.parse_args()


if __name__ == "__main__":
    mspattern = re.compile('\[(\d+)\.?(\d*)\]')  # matches either a page or line milestone
    kwargs = vars(args)
    folder = kwargs['folder']
    chngct = 0

    print(f"Updating documents in {folder}:")
    docs = os.listdir(folder)
    docs.sort()
    for docnm in docs:
        if '.docx' not in docnm:
            continue
        print(docnm)
        docpath = os.path.join(folder, docnm)
        doc = Document(docpath)
        p = doc.paragraphs[0]
        change = False
        for r in p.runs:
            mtchs = re.search(mspattern, r.text)
            if mtchs:
                pg = mtchs.group(1)
                ln = mtchs.group(2)
                print(pg, ln)
                if ln == "" or int(ln) > 2:
                    print("problem")
                elif int(ln) == 2:
                    # If first milestone is line 2, then insert page and line 1 Milestone at beginning of doc
                    new_run_element = p._element._new_r()
                    p.runs[0]._element.addprevious(new_run_element)
                    new_run = Run(new_run_element, p.runs[0]._parent)
                    new_run.text = r.text.replace('.2', '.1')
                    new_run.style = r.style
                    new_run.font.name = "Microsoft Himalaya (Body CS)"
                    new_run.font.size = r.style.font.size
                    # Insert Page Milestone
                    new_run_element = p._element._new_r()
                    p.runs[0]._element.addprevious(new_run_element)
                    new_run = Run(new_run_element, p.runs[0]._parent)
                    new_run.text = r.text.replace('.2', '')
                    new_run.style = "Page Number Print Edition"
                    new_run.font.name = "Microsoft Himalaya (Body CS)"
                    new_run.font.size = r.style.font.size

                    change = True
                break

        if change:
            newdocpath = os.path.join(folder, f"{docnm}-fixed.docx")
            doc.save(newdocpath)
        # for p in doc.paragraphs:
        #     for r in p.runs:
        #         mtchs = re.search(mspattern, r.text)
        #         if mtchs:
        #             pgnum = int(mtchs.group(1))
        #             lnstr = mtchs.group(2)
        #             if pgnum >= stnum and (endnum is None or pgnum <= endnum):
        #                 r.text = "[{}{}]".format(pgnum + delta, lnstr)
        #                 chngct += 1
        #
        # doc.save(docpath)
        # print("{} milestones changed!".format(chngct))

